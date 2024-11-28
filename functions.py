import openai
import os
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from re import findall
import json

def custom_input(msg):
    print(msg)
    lines = []
    while True:
        line = input()
        if line == "END":  # End of input indicator
            break
        lines.append(line)
    return "\n".join(lines)

# initialize OpenAI client
def initialize_openai_client(api_key_env="OPENAI_API_KEY"):
    """
    Initialize the OpenAI client using the API key from the environment variable.
    """
    api_key = os.getenv(api_key_env)
    if not api_key:
        raise ValueError("API key not found. Please set the API key in the environment variable.")
    return openai.OpenAI(api_key=api_key)


class ChatGPTSessionManager:
    def __init__(self, client, history_file='chat_history.json'):
        self.client = client
        self.history_file = history_file
        self.history = []

    # add a message to the conversation history
    def add_to_history(self, role, content):
        self.history.append({"role": role, "content": content})

    # send a message to ChatGPT and get a response
    def chat(self, message):
        self.add_to_history("user", message)
        response = self.client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=self.history
        )
        chat_response = response.choices[0].message.content
        self.add_to_history("assistant", chat_response)
        return chat_response


    # save conversation history to a file
    def save_history(self):
        with open(self.history_file, 'w') as f:
            json.dump(self.history, f, indent=4)


# read Python code from a file
def read_python_code(file_path):
    """
    Reads a Python code file and returns its content as a string.
    """
    if not file_path.endswith(".py"):
        raise ValueError(f"The file '{file_path}' is not a Python (.py) file.")
    try:
        with open(file_path, "r") as file:
            return file.read()
    except FileNotFoundError:
        raise FileNotFoundError(f"The file '{file_path}' was not found.")


# save Markdown-formatted documentation to Word
def save_markdown_to_word(markdown_text, output_file="documentation.docx"):
    """
    Converts Markdown-formatted text to a Word document with improved formatting that supports inline styles.
    """
    document = Document()

    # add a centered title with proper styling
    title = document.add_heading("Generated Documentation", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.style.font.size = Pt(14)
    title.style.font.bold = True

    # process Markdown lines
    lines = markdown_text.split("\n")
    in_code_block = False

    for line in lines:
        line = line.strip()

        if line.startswith("# "):  # Level 1 heading
            document.add_heading(line[2:].strip(), level=1)

        elif line.startswith("## "):  # Level 2 heading
            document.add_heading(line[3:].strip(), level=2)

        elif line.startswith("### "):  # Level 3 heading
            document.add_heading(line[4:].strip(), level=3)

        elif line.startswith("```"):  # Toggle code block
            in_code_block = not in_code_block

        elif in_code_block:  # Handle content inside code blocks
            code_paragraph = document.add_paragraph(line)
            code_paragraph.style.font.name = "Courier New"
            code_paragraph.style.font.size = Pt(10)

        else:  # Regular paragraphs with inline formatting
            paragraph = document.add_paragraph()

            # Inline formatting: handle **bold**, bullet points, and regular text
            elements = findall(r"(\*\*.*?\*\*|[-•] .*|.+)", line)

            for element in elements:
                if element.startswith("**") and element.endswith("**"):  # Bold inline text
                    run = paragraph.add_run(element.strip("**"))
                    run.bold = True
                elif element.startswith("- ") or element.startswith("• "):  # Bullet points
                    paragraph.style = "List Bullet"
                    paragraph.add_run(element.strip("- ").strip("• ").strip())
                else:  # Regular text
                    paragraph.add_run(element)

    # Add footer
    section = document.sections[-1]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_paragraph.style.font.size = Pt(8)

    # Save the document
    document.save(output_file)
    print(f"Documentation saved to {output_file}")


# Generate documentation with GPT-3.5 Turbo in English
def generate_documentation(client, code_snippets, style, additional_requests=None):
    """
    Generate documentation in English using GPT-3.5 Turbo.
    """
    combined_code = "\n\n".join(code_snippets)
    system_prompt = f"""
You are an elite-level Python programmer, software architect, and technical writer with extensive experience in creating high-quality technical documentation for research, conferences, and publications. Your task is to analyze the provided Python code and craft a creative, detailed, and logically structured documentation that is suitable for publication in a professional or academic setting. The documentation must reflect exceptional quality, ensuring clarity, precision, and depth.

The documentation should adhere to the following structure:

1. **Abstract**:
   - Provide a concise summary of the code's purpose, scope, and significance.
   - Highlight the key functionalities and the problem the code addresses.
   - BE CREATIVE! Make the abstract engaging and informative.
2. **Introduction**:
   - Describe the context and motivation behind the code.
   - Explain why the code is important, its real-world applications, and its potential impact.
   - Connect the code's relevance to broader fields or specific challenges.
   - Introduce the key concepts, tools, or technologies used in the code.
   -BE CREATIVE! Engage the reader from the beginning.

3. **Methodology**:
   - Offer an in-depth explanation of the code's structure, including its classes, functions, and algorithms.
   - Provide a high-level overview of the logic and workflow of the program.
   - Elaborate on how the components interact with one another to achieve the program's objectives.
   - Avoid bullet points or lists; instead, craft a narrative that flows logically.
4. **Technical Details**:
   - Dive deep into the implementation details:
     - Explain key algorithms, design decisions, and their benefits.
     - Describe inputs, outputs, and the role of important parameters.
   - Provide technical insights that demonstrate your mastery of the subject matter.
   - Avoid bullet points or lists; instead, craft a narrative that flows logically.

5. **Use Cases and Examples**:
   - Present real-world or hypothetical use cases for the code.
   - Show practical examples with sample inputs and expected outputs.
   - Illustrate how the code can be extended or modified for additional purposes.

6. **Discussion**:
   - Reflect on the strengths and limitations of the code.
   - Explore potential optimizations, scalability, or integration opportunities.
   - Suggest areas for future work or research.
- Avoid bullet points or lists; instead, craft a narrative that flows logically.
7. **Conclusion**:
   - Summarize the key takeaways from the documentation.
   - Reinforce the importance and contribution of the code.
   - Provide a final remark that leaves a lasting impression on the reader.

8. **Appendix (if applicable)**:
   - Include additional resources, references, file structure, or instructions for reproducing results.

**Requirements for the Writing Style**:
- The documentation must flow as a cohesive narrative, avoiding lists or bullet points wherever possible.
- Use precise, professional, and engaging language that demonstrates expertise while remaining approachable to a broad audience.
- The tone should be academic yet creative, ensuring the documentation reads like a conference paper or technical article.
- Explain concepts clearly but thoroughly, ensuring readers of varying technical backgrounds can follow along. Be creative in your explanations. Explain functions, and how to use the program.

**Formatting Guidelines**:
- Write in Markdown for easy conversion to other formats.
- Include clear headings and subheadings for organization.
- Use bold text, italics, and code blocks as necessary to emphasize key points or sections of the code.
"""
    if additional_requests:
        custom_requests = "\n".join([f"- {req}" for req in additional_requests])
        system_prompt += f"\nAdditionally, address the following user requests:\n{custom_requests}\n"

    session_manager = ChatGPTSessionManager(client)
    session_manager.add_to_history("system", system_prompt)
    response = session_manager.chat(combined_code)
    return response

# Translate English text to Hungarian
def translate_to_hungarian(client, english_text):
    """
    Translate English text to Hungarian using GPT-3.5 Turbo.
    """
    session_manager = ChatGPTSessionManager(client)

    prompt = f"""
Translate the following English text into fluent Hungarian while preserving its professional tone and logical structure:

{english_text}
"""
    response = session_manager.chat(prompt)
    return response.choices[0].message.content.strip()
